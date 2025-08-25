from flask import Flask, render_template, session, request, jsonify, redirect, url_for, session, send_file
from werkzeug.utils import secure_filename
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem
from reportlab.lib.units import inch
import os
import json
from dotenv import load_dotenv

load_dotenv()

# ------------------ Flask Setup ------------------
app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY")
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ------------------ Gemini Setup (Hardcoded as requested) ------------------
genai.configure(api_key=os.getenv("GENAI_API_KEY"))
# Model for everything
model = genai.GenerativeModel(
    "gemini-1.5-flash",
    system_instruction="""
You are a friendly mentor for beginners.  
Your mission is to guide people step by step towards their dreams in an **interactive and engaging conversation**.

⚡ Formatting Rules:
- Use <h1> for the main session title.  
- Use <h2> for section titles (e.g., Job Roles in AI).  
- Use <h3> for sub-sections (e.g., Machine Learning Engineer).  
- Use <ul><li> for bullet points.  
- Use <b> for highlighting important terms, resources, or key ideas.  
- Keep everything left-aligned, clean, and beginner-friendly.  

⚡ Conversation Flow:
1. Always begin by warmly asking about the user’s **interests, dream career, or goals**.  
   Example: “What’s something you dream of doing in the future? Or do you just want me to show you different exciting fields to choose from?”  

2. If the user already has a field of interest:  
   - Start by explaining what that field is.  
   - Then show **real-world applications**.  
   - Then explain **job roles** in that field.  
   - Finally, provide a **roadmap with resources, project ideas, and steps**.  
   - At the end of each section, ask: “Would you like me to explain this next?”  

3. If the user does NOT know their interest:  
   - Start introducing popular and beginner-friendly fields one by one (e.g., AI, Web Development, Cybersecurity, Design, etc.).  
   - For each field:  
     <ul>
     <li>Give a **short simple explanation**.</li>  
     <li>Show **real-world applications** and how it impacts daily life.</li>  
     <li>Ask the user: “Does this spark your interest? Should I tell you more?”</li>  
     </ul>  
   - Continue until the user feels excited about a field.  
   - Once they choose, follow the same flow as above (Applications → Job Roles → Roadmap).  

4. Always keep the tone **supportive, empathetic, and motivating**. Make the user feel guided by a real mentor, not just a bot.  

⚡ Important:
- Do NOT dump everything at once.  
- Keep conversations ongoing by asking questions and waiting for the user’s input.  
- The main goal is to **either guide them in their chosen field OR help them discover one**. 
- Dont leave Line where ever line is to be left instead start writing from that line onewards
"""
)

# ------------------ User Store ------------------
USERS_FILE = "users.json"

def load_users():
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, "r") as f:
            return json.load(f)
    return {}

def save_users(users):
    with open(USERS_FILE, "w") as f:
        json.dump(users, f, indent=2)


# Career chat (kept as-is but isolated per session)
def get_chat():
    if "chat_session" not in session:
        # start a new chat history for this user session
        session["chat_session"] = []
    return session["chat_session"]

def append_chat(role, content):
    chat_hist = get_chat()
    chat_hist.append({"role": role, "content": content})
    session["chat_session"] = chat_hist

# ------------------ Helpers ------------------

def extract_text_from_pdf(file_stream) -> str:
    """Extract text from an uploaded PDF (file stream)."""
    reader = PdfReader(file_stream)
    text = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        text.append(txt)
    return "\n".join(text).strip()

def extract_text_from_docx(file_stream) -> str:
    """Extract text from an uploaded DOCX (file stream)."""
    # python-docx expects a path-like or file-like, BytesIO is fine
    doc = Document(file_stream)
    return "\n".join(p.text for p in doc.paragraphs).strip()

def gemini_structured_resume(conversation_messages):
    """
    Ask Gemini to convert chat conversation into a structured JSON resume.
    Returns dict with keys: contact, summary, education(list), experience(list), skills(list), projects(list)
    """
    convo_text = "\n".join(f"{m['role']}: {m['content']}" for m in conversation_messages)
    prompt = f"""
Convert the following conversation into a concise, structured resume JSON.

Rules:
- Return ONLY valid JSON (no prose).
- Keys: contact (string), summary (string), education (list of strings),
        experience (list of strings), skills (list of strings), projects (list of strings)

Conversation:
{convo_text}
"""
    response = model.generate_content(prompt)
    raw = response.text.strip()
    try:
        data = json.loads(raw)
        # sanity defaults
        data.setdefault("contact", "")
        data.setdefault("summary", "")
        data.setdefault("education", [])
        data.setdefault("experience", [])
        data.setdefault("skills", [])
        data.setdefault("projects", [])
        return data
    except Exception:
        # Fallback: produce a minimal dict using plain text blob
        return {
            "contact": "",
            "summary": raw[:800],
            "education": [],
            "experience": [],
            "skills": [],
            "projects": []
        }

def build_pdf_from_resume(resume_dict) -> BytesIO:
    """
    Build a professional-looking PDF from structured resume data using ReportLab.
    Returns BytesIO buffer.
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=36, rightMargin=36, topMargin=54, bottomMargin=36
    )
    styles = getSampleStyleSheet()

    # Custom header styles
    h_name = ParagraphStyle(
        "NameHeader", parent=styles["Title"], fontSize=20, leading=24, spaceAfter=8
    )
    h_sec = ParagraphStyle(
        "SectionHeader", parent=styles["Heading2"], textColor=colors.HexColor("#333333"),
        spaceBefore=10, spaceAfter=6, underlineWidth=0.5
    )
    body = styles["Normal"]
    body.spaceAfter = 2

    story = []

    # Contact / Name line
    contact = resume_dict.get("contact", "").strip()
    if contact:
        story.append(Paragraph(contact, h_name))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#999999")))
        story.append(Spacer(1, 0.15*inch))

    # Summary
    summary = resume_dict.get("summary", "").strip()
    if summary:
        story.append(Paragraph("Summary", h_sec))
        story.append(Paragraph(summary, body))
        story.append(Spacer(1, 0.1*inch))

    # Education
    education = resume_dict.get("education", [])
    if education:
        story.append(Paragraph("Education", h_sec))
        edu_list = ListFlowable(
            [ListItem(Paragraph(e, body), leftIndent=12) for e in education],
            bulletType="bullet", start="circle"
        )
        story.append(edu_list)
        story.append(Spacer(1, 0.08*inch))

    # Experience
    experience = resume_dict.get("experience", [])
    if experience:
        story.append(Paragraph("Experience", h_sec))
        exp_list = ListFlowable(
            [ListItem(Paragraph(x, body), leftIndent=12) for x in experience],
            bulletType="bullet", start="circle"
        )
        story.append(exp_list)
        story.append(Spacer(1, 0.08*inch))

    # Skills
    skills = resume_dict.get("skills", [])
    if skills:
        story.append(Paragraph("Skills", h_sec))
        # render skills as comma-separated paragraph
        story.append(Paragraph(", ".join(skills), body))
        story.append(Spacer(1, 0.08*inch))

    # Projects
    projects = resume_dict.get("projects", [])
    if projects:
        story.append(Paragraph("Projects", h_sec))
        proj_list = ListFlowable(
            [ListItem(Paragraph(p, body), leftIndent=12) for p in projects],
            bulletType="bullet", start="circle"
        )
        story.append(proj_list)

    doc.build(story)
    buf.seek(0)
    return buf

def build_docx_from_resume(resume_dict) -> BytesIO:
    """
    Build a DOCX resume using python-docx.
    """
    doc = Document()

    contact = resume_dict.get("contact", "").strip()
    if contact:
        h = doc.add_heading(contact, level=0)
    doc.add_paragraph("")  # spacer

    summary = resume_dict.get("summary", "").strip()
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    education = resume_dict.get("education", [])
    if education:
        doc.add_heading("Education", level=1)
        for e in education:
            doc.add_paragraph(e, style="List Bullet")

    experience = resume_dict.get("experience", [])
    if experience:
        doc.add_heading("Experience", level=1)
        for x in experience:
            doc.add_paragraph(x, style="List Bullet")

    skills = resume_dict.get("skills", [])
    if skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(skills))

    projects = resume_dict.get("projects", [])
    if projects:
        doc.add_heading("Projects", level=1)
        for p in projects:
            doc.add_paragraph(p, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ------------------ Routes (Pages) ------------------

@app.route("/")
def home():
    return redirect(url_for("login"))

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")
        users = load_users()

        # simple check
        if email in users and users[email]["password"] == password:
            session["user"] = email
            return redirect(url_for("main"))
        else:
            return "Invalid credentials", 401
    return render_template("login.html")



@app.route('/logout', methods=['POST'])
def logout():
    session.clear()  
    return redirect(url_for('login')) 

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        first = request.form.get("firstName")
        last = request.form.get("lastName")
        email = request.form.get("email")
        password = request.form.get("password")

        users = load_users()
        users[email] = {"first": first, "last": last, "password": password}
        save_users(users)

        return redirect(url_for("login"))
    return render_template("register.html")

@app.route("/main")
def main():
    if "user" not in session:
        return redirect(url_for("login"))

    users = load_users()
    user_email = session["user"]
    user_data = users.get(user_email, {})
    full_name = f"{user_data.get('first', '')} {user_data.get('last', '')}".strip()

    return render_template("main.html",
                           username=full_name,
                           email=user_email)

@app.route("/career")
def career():
    if "user" not in session:
        return redirect(url_for("login"))
    return render_template("career.html")

@app.route("/analyzer")
def analyzer():
    if "user" not in session:
        return redirect(url_for("login"))
    return render_template("analyzer.html")

@app.route("/builder", methods=["POST"])
def resume_builder_chat():
    data = request.get_json(silent=True) or {}
    user_msg = (data.get("message") or "").strip()

    if not user_msg:
        return jsonify({"reply": "⚠ Please type something."})

    # Session state
    step = session.get("resume_step", 0)              # 0..len(compulsory_qs)
    answers = session.get("resume_answers", {})       # dict

    # Commands
    cmd = user_msg.lower()
    if cmd == "/restart_resume":
        for k in ["resume_step", "resume_answers", "custom_qs", "custom_step", "final_resume"]:
            session.pop(k, None)
        session["resume_step"] = 0
        session["resume_answers"] = {}
        return jsonify({"reply": "🔄 Restarting… Let's start again. What is your full name?"})

    if cmd == "/resume":
        if "final_resume" in session:
            return jsonify({"reply": "✅ Your resume is ready! Type /download to get your PDF."})
        # Not ready — tell user what’s missing
        pieces = []
        compulsory_needed = max(0, 5 - step)
        if compulsory_needed:
            pieces.append(f"{compulsory_needed} compulsory question(s) remaining.")
        custom_qs = session.get("custom_qs", [])
        custom_step = session.get("custom_step", 0)
        if custom_qs and custom_step < len(custom_qs):
            pieces.append(f"{len(custom_qs) - custom_step} follow-up question(s) remaining.")
        if not pieces:
            pieces.append("Say anything to continue and I’ll finalize it.")
        return jsonify({"reply": "⚠ Your resume isn’t ready yet. " + " ".join(pieces)})

    if cmd == "/download":
        if "final_resume" not in session:
            return jsonify({"reply": "⚠ No resume generated yet. Type /resume when you're done answering questions."})
        # IMPORTANT: return a link, not the file, so the fetch() in your chat UI doesn't break
        return jsonify({
            "reply": "⬇ Click to download your PDF: <a href='/download_resume' target='_blank'>Download Resume</a>"
        })

    # Compulsory questions
    compulsory_qs = [
        "What is your full name?",
        "What is your email?",
        "What is your phone number?",
        "What is your field of interest?",
        "What are your top skills?"
    ]

    # Still collecting compulsory questions?
    if step < len(compulsory_qs):
        # Save the answer for the current expected question (including the first one)
        answers[f"Q{step+1}"] = user_msg
        session["resume_answers"] = answers
        session["resume_step"] = step + 1

        # Ask next compulsory question if any
        if step + 1 < len(compulsory_qs):
            return jsonify({"reply": compulsory_qs[step + 1]})
        else:
            # Generate custom questions based on skills
            skills = answers.get("Q5", "")
            prompt = (
                "Generate 5 short, distinct questions (one per line) to ask a candidate "
                f"with skills: {skills}, to make a strong resume. Do not number them."
            )
            try:
                gen = model.generate_content(prompt)
                lines = (gen.text or "").splitlines()
                # Clean bullet/number prefixes
                import re
                custom_qs = []
                for ln in lines:
                    ln = ln.strip()
                    if not ln:
                        continue
                    ln = re.sub(r'^\s*(?:\d+[\).\s-]|[-•]\s)', '', ln)
                    if ln:
                        custom_qs.append(ln)
                # Fallback if model returns nothing useful
                if len(custom_qs) < 3:
                    raise ValueError("Too few generated questions")
            except Exception:
                custom_qs = [
                    "Tell me about a project where you used those skills.",
                    "What achievements are you most proud of?",
                    "Any certifications or courses completed?",
                    "What tools/technologies do you use most?",
                    "What kind of roles are you targeting next?"
                ]
            session["custom_qs"] = custom_qs
            session["custom_step"] = 0
            return jsonify({"reply": custom_qs[0]})

    # Custom, skill-based questions
    custom_qs = session.get("custom_qs", [])
    custom_step = session.get("custom_step", 0)

    if custom_qs and custom_step < len(custom_qs):
        # Save answer to the current custom question
        answers[f"CustomQ{custom_step+1}"] = user_msg
        session["resume_answers"] = answers
        custom_step += 1
        session["custom_step"] = custom_step

        # More custom questions remain?
        if custom_step < len(custom_qs):
            return jsonify({"reply": custom_qs[custom_step]})
        else:
            # Build resume
            try:
                prompt = (
                    "Create a professional, concise resume in clean text (not code) "
                    "based on the following details:\n"
                    f"{json.dumps(session['resume_answers'], ensure_ascii=False, indent=2)}"
                )
                resume_draft = model.generate_content(prompt).text or ""
            except Exception:
                resume_draft = "Resume draft could not be generated automatically. Please try again."

            session["final_resume"] = resume_draft.strip()
            return jsonify({
                "reply": "✅ Your resume is ready! Type /download to get the PDF or /restart_resume to start again.",
                "resume": session["final_resume"]
            })

    # Safety fallback
    session["custom_qs"] = []
    session["custom_step"] = 0
    return jsonify({"reply": "I’m ready to continue. Tell me more about your projects or type /resume to generate your resume."})
# ------------------ API: Career Chat (kept) ------------------

@app.route("/chat", methods=["POST"])
def chat_with_ai():
    user_input = (request.json or {}).get("message", "").strip()
    if not user_input:
        return jsonify({"reply": "⚠️ Please type something."})
    try:
        # keep per-session chat memory using messages array (simple)
        append_chat("user", user_input)
        # generate assistant response using full session history
        hist = get_chat()
        # build a single prompt from history (simple approach)
        history_text = "\n".join(f"{m['role']}: {m['content']}" for m in hist)
        response = model.generate_content(history_text + f"\nassistant: ")
        reply = response.text
        append_chat("assistant", reply)
        return jsonify({"reply": reply})
    except Exception as e:
        return jsonify({"reply": f"Error: {str(e)}"})

# ------------------ API: Resume Analyzer ------------------

@app.route("/analyze_resume", methods=["POST"])
def analyze_resume():
    try:
        # --- Get Job Description ---
        job_desc = (request.form.get("jobDescription") or "").strip()

        # --- Get Resume Text (pasted) ---
        resume_text = (request.form.get("resumeText") or "").strip()

        # --- If File Uploaded, Use That Instead ---
        if "resume" in request.files and request.files["resume"].filename:
            f = request.files["resume"]
            filename = secure_filename(f.filename)
            ext = (os.path.splitext(filename)[1] or "").lower()

            if ext == ".pdf":
                resume_text = extract_text_from_pdf(f.stream)
            elif ext in (".docx", ".doc"):
                resume_text = extract_text_from_docx(f.stream)
            else:
                return jsonify({"error": "Only PDF or DOCX files are supported"})

        # --- Validate Inputs ---
        if not job_desc or not resume_text:
            return jsonify({"error": "⚠️ Resume text/file and job description are required"})

        # --- Gemini ATS Prompt ---
        prompt = f"""
        You are an advanced Applicant Tracking System (ATS) evaluator.

        TASKS:
        1. Calculate an **ATS Compatibility Score** (0-100) of the resume against the job description.
        2. Provide **detailed suggestions** on how the resume can be improved to match the job description better (skills, keywords, formatting, etc.).
        3. Suggest **job opportunities / roles** that the candidate could apply for, based on their resume content and skillset.

        Resume:
        {resume_text}

        Job Description:
        {job_desc}

        Respond in the following structured format:
        - ATS Score: <score out of 100>
        - Suggestions: <bullet points>
        - Job Opportunities: <list of possible job roles>
        """

        # --- Call Gemini ---
        response = model.generate_content(prompt)

        return jsonify({"analysis": response.text})

    except Exception as e:
        return jsonify({"error": str(e)})



# ------------------ API: Resume Builder Chat + Exports ------------------

@app.route("/resume_builder", methods=["POST"])
def resume_builder():
    """
    Chat endpoint for builder.html.
    - Stores conversation in session["conversation"]
    - If user types '/resume' in FE, they will call /generate_resume instead
    """
    msg = (request.json or {}).get("message", "").strip()
    if not msg:
        return jsonify({"reply": "⚠️ Please type something."})

    try:
        # log conversation for this session
        conversation = session.get("conversation", [])
        conversation.append({"role": "user", "content": msg})

        # Guide the user step-by-step
        guide = """
You are a Resume Builder AI.
- Ask for details step-by-step (name & contact, summary, education, experience, skills, projects).
- After each user reply, ask the next most relevant question.
- Keep answers short. Do NOT dump the full resume at once.
"""
        history_text = "\n".join(f"{m['role']}: {m['content']}" for m in conversation)
        prompt = guide + "\n" + history_text + "\nassistant:"
        response = model.generate_content(prompt)
        reply = response.text

        conversation.append({"role": "ai", "content": reply})
        session["conversation"] = conversation
        return jsonify({"reply": reply})
    except Exception as e:
        return jsonify({"reply": f"Error: {str(e)}"})

@app.route("/generate_resume", methods=["POST"])
def generate_resume():
    """
    Generates a professional PDF from the current conversation.
    """
    try:
        conversation = session.get("conversation", [])
        if not conversation:
            return "No conversation yet.", 400

        resume_json = gemini_structured_resume(conversation)
        session["resume_json"] = resume_json

        pdf_buf = build_pdf_from_resume(resume_json)
        return send_file(pdf_buf, as_attachment=True,
                         download_name="My_Resume.pdf",
                         mimetype="application/pdf")
    except Exception as e:
        return f"Error generating PDF: {str(e)}", 500

@app.route("/generate_resume_docx", methods=["POST"])
def generate_resume_docx():
    """
    Generates a DOCX resume from session resume_json (or conversation if missing).
    """
    try:
        resume_json = session.get("resume_json")
        if not resume_json:
            conversation = session.get("conversation", [])
            if not conversation:
                return "No data to build resume.", 400
            resume_json = gemini_structured_resume(conversation)
            session["resume_json"] = resume_json

        docx_buf = build_docx_from_resume(resume_json)
        return send_file(docx_buf, as_attachment=True,
                         download_name="My_Resume.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return f"Error generating DOCX: {str(e)}", 500

# ------------------ Run ------------------
if __name__ == "__main__":
    app.run(debug=True) 